import os
import re
import sys

# Attempt to import docx, instructions provided if missing
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    print("Error: 'python-docx' library is not installed.")
    print("Please run: pip install python-docx")
    sys.exit(1)

def set_cell_background(cell, hex_color):
    """Set the background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, text):
    """Add a shaded box containing code or math text."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    
    # Set thin border
    tcPr = cell._tc.get_or_add_tcPr()
    borders_elm = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="0D9488"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders_elm)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(15, 23, 42)

def add_styled_paragraph(doc, text, style_name=None):
    """Add a paragraph and parse basic markdown inline styling (**bold**, *italic*)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    # Tokenize text by markdown bold and italics tags
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = p.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(13, 148, 136)
        else:
            p.add_run(part)
            
    return p

def add_styled_bullet(doc, text):
    """Add a bullet point parsing bold/italic tags."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = p.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(13, 148, 136)
        else:
            p.add_run(part)

def build_docx_report():
    print("Initializing document compiler...")
    doc = docx.Document()
    
    # Page Setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    styles = doc.styles
    
    # Configure Normal/Body Text Style
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(10.5)
    font_normal.color.rgb = RGBColor(51, 65, 85) # Slate 700
    
    # Document Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("GENMED HUB™")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(32)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("System Architecture, Custom Algorithms & API Specifications Manual")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
    
    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(100)
    run_desc = p_desc.add_run("A Comprehensive Guide to the Inner Workings of the Indian Generic Medicine Platform")
    run_desc.font.name = 'Arial'
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    # Metadata info block on cover
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.line_spacing = 1.3
    p_meta.paragraph_format.space_after = Pt(6)
    
    runs_meta = [
        ("Platform Type: ", True), ("React 19 + Node.js REST Application\n", False),
        ("Database Layer: ", True), ("Flat JSON Database (520+ Medicines, 515 Diseases)\n", False),
        ("Core Models: ", True), ("Haversine Spatial Proximity, Simple Linear Regression with Seasonality\n", False),
        ("Generated Date: ", True), ("June 2026\n", False),
        ("Author: ", True), ("Antigravity AI Pair-Programming Assistant", False)
    ]
    
    for text, is_bold in runs_meta:
        run = p_meta.add_run(text)
        run.bold = is_bold
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(100, 116, 139)
        
    doc.add_page_break()
    
    # Ordered markdown files list to build doc content
    docs_dir = "documentation"
    markdown_files = [
        "01_acknowledgements.md",
        "02_problem_definition_and_existing_system.md",
        "03_proposed_system_objectives_and_scope.md",
        "04_requirement_gathering_and_specifications.md",
        "05_technology_stack_and_system_architecture.md",
        "06_module_description_and_functions.md",
        "07_mathematical_algorithms_and_heuristics.md",
        "08_implementation_strategies_and_testing.md",
        "09_limitations_and_future_scope.md"
    ]
    
    for index, filename in enumerate(markdown_files):
        filepath = os.path.join(docs_dir, filename)
        if not os.path.exists(filepath):
            print(f"Warning: File {filepath} not found. Skipping...")
            continue
            
        print(f"Parsing: {filename}...")
        
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Parse lines
        lines = content.split('\n')
        
        in_code_block = False
        code_lines = []
        
        in_table = False
        table_rows = []
        
        for line in lines:
            stripped = line.strip()
            
            # 1. Code block handling
            if stripped.startswith('```'):
                if in_code_block:
                    # End of code block
                    add_code_block(doc, '\n'.join(code_lines))
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue
                
            if in_code_block:
                code_lines.append(line)
                continue
                
            # 2. Table handling
            if stripped.startswith('|'):
                in_table = True
                # Parse columns
                cols = [col.strip() for col in line.split('|')[1:-1]]
                table_rows.append(cols)
                continue
            elif in_table:
                # If table just ended
                in_table = False
                # Filter out separator rows e.g. |---|---|
                cleaned_rows = [r for r in table_rows if not all(re.match(r'^-+$', c) for c in r)]
                table_rows = []
                
                if cleaned_rows:
                    num_rows = len(cleaned_rows)
                    num_cols = len(cleaned_rows[0])
                    
                    docx_table = doc.add_table(rows=num_rows, cols=num_cols)
                    docx_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    docx_table.style = 'Table Grid'
                    
                    for r_idx, row_data in enumerate(cleaned_rows):
                        for c_idx, cell_value in enumerate(row_data):
                            if c_idx < num_cols:
                                cell = docx_table.cell(r_idx, c_idx)
                                cell.text = cell_value
                                
                                # Set header row styles
                                if r_idx == 0:
                                    set_cell_background(cell, "0F172A") # Slate 900
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                            run.font.size = Pt(9)
                                else:
                                    # Alternating row background
                                    if r_idx % 2 == 1:
                                        set_cell_background(cell, "F8FAFC")
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(8.5)
                                            run.font.color.rgb = RGBColor(51, 65, 85)
                                            
                    # Add paragraph spacing after table
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
                    
            if not stripped:
                continue
                
            # 3. Headings handling
            if stripped.startswith('# '):
                # Ignore top main title from file if it is just section title
                title_text = stripped[2:]
                h = doc.add_heading(level=1)
                h.paragraph_format.space_before = Pt(16)
                h.paragraph_format.space_after = Pt(8)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(title_text)
                run.font.name = 'Arial'
                run.bold = True
                run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
            elif stripped.startswith('## '):
                h_text = stripped[3:]
                h = doc.add_heading(level=2)
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(6)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(h_text)
                run.font.name = 'Arial'
                run.bold = True
                run.font.color.rgb = RGBColor(13, 148, 136) # Teal 700
            elif stripped.startswith('### '):
                h_text = stripped[4:]
                h = doc.add_heading(level=3)
                h.paragraph_format.space_before = Pt(10)
                h.paragraph_format.space_after = Pt(4)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(h_text)
                run.font.name = 'Arial'
                run.bold = True
                run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
            elif stripped.startswith('#### '):
                h_text = stripped[5:]
                h = doc.add_heading(level=4)
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after = Pt(4)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(h_text)
                run.font.name = 'Arial'
                run.bold = True
                run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
            
            # 4. Bullet list items
            elif stripped.startswith('- ') or stripped.startswith('* '):
                item_text = stripped[2:]
                add_styled_bullet(doc, item_text)
                
            elif re.match(r'^\d+\.\s', stripped):
                # Ordered lists - just render as regular styled text with numbering
                add_styled_paragraph(doc, stripped)
                
            # 5. Normal text paragraphs
            else:
                # Skip mermaid diagrams text tags
                if stripped.startswith('graph') or stripped.startswith('subgraph') or '-->' in stripped or '---' in stripped:
                    continue
                add_styled_paragraph(doc, stripped)
                
        # Add Page Break between major sections except the last one
        if index < len(markdown_files) - 1:
            doc.add_page_break()
            
    output_filename = "GenMed_Hub_Project_Report.docx"
    doc.save(output_filename)
    print(f"Report compiled successfully: {output_filename}")

if __name__ == "__main__":
    build_docx_report()
